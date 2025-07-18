"""
Document processing module for legal documents.
Handles text extraction, cleaning, and chunking optimized for legal content.
"""

import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import textract
from bs4 import BeautifulSoup
import chardet

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy

from ..config.settings import Settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced document processor optimized for legal documents.
    Handles multiple formats with specialized legal text processing.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.settings = Settings()
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt', '.html', '.rtf'}
        
        # Configurações de chunking
        self.chunk_size = self.settings.CHUNK_SIZE
        self.chunk_overlap = self.settings.CHUNK_OVERLAP
        self.min_chunk_size = 100
        
        # Padrões jurídicos brasileiros
        self.legal_patterns = {
            'artigo': re.compile(r'art\.?\s*\d+[º°]?', re.IGNORECASE),
            'inciso': re.compile(r'inciso\s+[IVX]+', re.IGNORECASE),
            'paragrafo': re.compile(r'§\s*\d+[º°]?', re.IGNORECASE),
            'lei': re.compile(r'lei\s+n[º°]?\s*[\d\.]+', re.IGNORECASE),
            'decreto': re.compile(r'decreto\s+n[º°]?\s*[\d\.]+', re.IGNORECASE),
            'processo': re.compile(r'\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}', re.IGNORECASE),
            'cnj': re.compile(r'\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}', re.IGNORECASE)
        }
        
        # Inicializar NLP
        self._init_nlp()
        
        logger.info("DocumentProcessor inicializado")
    
    def _init_nlp(self):
        """Initialize NLP libraries."""
        try:
            # Download NLTK data if needed
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            
            # Load spaCy model for Portuguese
            try:
                self.nlp = spacy.load("pt_core_news_sm")
            except OSError:
                logger.warning("Modelo spaCy pt_core_news_sm não encontrado. Usando processamento básico.")
                self.nlp = None
            
            # Portuguese stopwords
            try:
                self.stopwords = set(stopwords.words('portuguese'))
            except:
                self.stopwords = set()
                
        except Exception as e:
            logger.warning(f"Erro ao inicializar NLP: {str(e)}")
            self.nlp = None
            self.stopwords = set()
    
    async def process_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Process a document and extract structured information.
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata for the document
            
        Returns:
            Dictionary with processed document information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {file_path.suffix}")
        
        logger.info(f"Processando documento: {file_path.name}")
        
        try:
            # Extrair texto
            raw_text = await self._extract_text(file_path)
            
            # Limpar texto
            cleaned_text = self._clean_text(raw_text)
            
            # Extrair metadados do texto
            extracted_metadata = self._extract_metadata(cleaned_text)
            
            # Combinar metadados
            final_metadata = {**(metadata or {}), **extracted_metadata}
            
            # Criar chunks
            chunks = self._create_chunks(cleaned_text)
            
            # Calcular hash do documento
            doc_hash = self._calculate_hash(cleaned_text)
            
            result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_type': file_path.suffix.lower(),
                'document_hash': doc_hash,
                'raw_text': raw_text,
                'cleaned_text': cleaned_text,
                'chunks': chunks,
                'metadata': final_metadata,
                'stats': {
                    'total_chars': len(cleaned_text),
                    'total_words': len(cleaned_text.split()),
                    'total_chunks': len(chunks),
                    'avg_chunk_size': sum(len(chunk['text']) for chunk in chunks) / len(chunks) if chunks else 0
                }
            }
            
            logger.info(f"Documento processado: {len(chunks)} chunks gerados")
            return result
            
        except Exception as e:
            logger.error(f"Erro ao processar documento {file_path}: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type."""
        loop = asyncio.get_event_loop()
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return await loop.run_in_executor(None, self._extract_pdf_text, file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return await loop.run_in_executor(None, self._extract_docx_text, file_path)
            elif file_path.suffix.lower() == '.txt':
                return await loop.run_in_executor(None, self._extract_txt_text, file_path)
            elif file_path.suffix.lower() == '.html':
                return await loop.run_in_executor(None, self._extract_html_text, file_path)
            else:
                # Fallback usando textract
                return await loop.run_in_executor(None, self._extract_with_textract, file_path)
                
        except Exception as e:
            logger.error(f"Erro na extração de texto: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods."""
        text = ""
        
        # Método 1: pdfplumber (melhor para tabelas e layout)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text
                
        except Exception as e:
            logger.warning(f"pdfplumber falhou: {str(e)}")
        
        # Método 2: PyPDF2 (fallback)
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text
                
        except Exception as e:
            logger.warning(f"PyPDF2 falhou: {str(e)}")
        
        # Método 3: textract (último recurso)
        try:
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except Exception as e:
            logger.error(f"Todos os métodos de extração PDF falharam: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            # Fallback para textract
            try:
                text = textract.process(str(file_path)).decode('utf-8')
                return text
            except:
                logger.error(f"Erro ao extrair texto DOCX: {str(e)}")
                raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            # Detectar encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            # Ler com encoding detectado
            with open(file_path, 'r', encoding=encoding or 'utf-8') as file:
                return file.read()
                
        except Exception as e:
            # Tentar com UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except:
                # Último recurso: latin-1
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
    
    def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"Erro ao extrair texto HTML: {str(e)}")
            raise
    
    def _extract_with_textract(self, file_path: Path) -> str:
        """Extract text using textract as fallback."""
        try:
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except Exception as e:
            logger.error(f"textract falhou: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for legal documents."""
        if not text:
            return ""
        
        # Remover caracteres de controle
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Normalizar espaços em branco
        text = re.sub(r'\s+', ' ', text)
        
        # Remover linhas muito curtas (provavelmente headers/footers)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 10:  # Manter apenas linhas com conteúdo substancial
                cleaned_lines.append(line)
        
        text = '\n'.join(cleaned_lines)
        
        # Normalizar quebras de linha
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Corrigir hifenização
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        return text.strip()
    
    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        """Extract metadata from legal document text."""
        metadata = {}
        
        # Detectar tipo de documento
        doc_type = self._detect_document_type(text)
        metadata['document_type'] = doc_type
        
        # Extrair entidades jurídicas
        legal_entities = self._extract_legal_entities(text)
        metadata.update(legal_entities)
        
        # Estatísticas do texto
        metadata['language'] = 'pt'  # Assumindo português
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
        
        return metadata
    
    def _detect_document_type(self, text: str) -> str:
        """Detect the type of legal document."""
        text_lower = text.lower()
        
        # Padrões para diferentes tipos de documentos
        patterns = {
            'lei': ['lei federal', 'lei estadual', 'lei municipal', 'lei nº', 'lei n°'],
            'decreto': ['decreto nº', 'decreto n°', 'decreto federal', 'decreto estadual'],
            'sentenca': ['sentença', 'julgo procedente', 'julgo improcedente', 'dispositivo'],
            'acordao': ['acórdão', 'tribunal', 'relator', 'revisor'],
            'peticao': ['petição inicial', 'requer', 'requerente', 'requerido'],
            'contrato': ['contrato', 'contratante', 'contratado', 'cláusula'],
            'parecer': ['parecer jurídico', 'parecer técnico', 'conclusão'],
            'ata': ['ata da reunião', 'ata de assembléia', 'presentes']
        }
        
        for doc_type, keywords in patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                return doc_type
        
        return 'documento_juridico'
    
    def _extract_legal_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract legal entities from text."""
        entities = {
            'leis': [],
            'decretos': [],
            'artigos': [],
            'processos': [],
            'paragrafos': [],
            'incisos': []
        }
        
        # Extrair usando padrões regex
        for entity_type, pattern in self.legal_patterns.items():
            matches = pattern.findall(text)
            if entity_type in ['lei', 'decreto']:
                entities[f'{entity_type}s'] = list(set(matches))
            elif entity_type == 'artigo':
                entities['artigos'] = list(set(matches))
            elif entity_type in ['processo', 'cnj']:
                entities['processos'].extend(matches)
            elif entity_type == 'paragrafo':
                entities['paragrafos'] = list(set(matches))
            elif entity_type == 'inciso':
                entities['incisos'] = list(set(matches))
        
        # Remover duplicatas
        entities['processos'] = list(set(entities['processos']))
        
        return entities
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create text chunks optimized for legal documents."""
        if not text:
            return []
        
        # Dividir por seções legais primeiro
        sections = self._split_by_legal_sections(text)
        
        chunks = []
        chunk_id = 0
        
        for section in sections:
            section_chunks = self._split_section_into_chunks(section, chunk_id)
            chunks.extend(section_chunks)
            chunk_id += len(section_chunks)
        
        return chunks
    
    def _split_by_legal_sections(self, text: str) -> List[str]:
        """Split text by legal sections (articles, paragraphs, etc.)."""
        # Padrões para divisão de seções
        section_patterns = [
            r'\n\s*Art\.?\s*\d+[º°]?',  # Artigos
            r'\n\s*§\s*\d+[º°]?',       # Parágrafos
            r'\n\s*Capítulo\s+[IVX]+',  # Capítulos
            r'\n\s*Seção\s+[IVX]+',     # Seções
            r'\n\s*Título\s+[IVX]+',    # Títulos
        ]
        
        # Encontrar pontos de divisão
        split_points = [0]
        
        for pattern in section_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                split_points.append(match.start())
        
        # Ordenar e remover duplicatas
        split_points = sorted(set(split_points))
        split_points.append(len(text))
        
        # Criar seções
        sections = []
        for i in range(len(split_points) - 1):
            start = split_points[i]
            end = split_points[i + 1]
            section = text[start:end].strip()
            
            if section and len(section) > self.min_chunk_size:
                sections.append(section)
        
        return sections if sections else [text]
    
    def _split_section_into_chunks(self, section: str, start_chunk_id: int) -> List[Dict[str, Any]]:
        """Split a section into chunks."""
        chunks = []
        
        # Se a seção é pequena, retornar como um chunk
        if len(section) <= self.chunk_size:
            chunks.append({
                'id': start_chunk_id,
                'text': section,
                'start_char': 0,
                'end_char': len(section),
                'metadata': self._extract_chunk_metadata(section)
            })
            return chunks
        
        # Dividir por sentenças
        sentences = self._split_into_sentences(section)
        
        current_chunk = ""
        current_start = 0
        chunk_id = start_chunk_id
        
        for sentence in sentences:
            # Verificar se adicionar a sentença excede o tamanho do chunk
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Salvar chunk atual
                chunks.append({
                    'id': chunk_id,
                    'text': current_chunk.strip(),
                    'start_char': current_start,
                    'end_char': current_start + len(current_chunk),
                    'metadata': self._extract_chunk_metadata(current_chunk)
                })
                
                # Iniciar novo chunk com overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + sentence
                current_start += len(current_chunk) - len(overlap_text)
                chunk_id += 1
            else:
                current_chunk += sentence
        
        # Adicionar último chunk se não estiver vazio
        if current_chunk.strip():
            chunks.append({
                'id': chunk_id,
                'text': current_chunk.strip(),
                'start_char': current_start,
                'end_char': current_start + len(current_chunk),
                'metadata': self._extract_chunk_metadata(current_chunk)
            })
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        try:
            if self.nlp:
                doc = self.nlp(text)
                return [sent.text + " " for sent in doc.sents]
            else:
                # Fallback para NLTK
                sentences = sent_tokenize(text, language='portuguese')
                return [sent + " " for sent in sentences]
        except:
            # Fallback simples
            sentences = re.split(r'[.!?]+\s+', text)
            return [sent + ". " for sent in sentences if sent.strip()]
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text for chunk continuity."""
        words = text.split()
        overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else words
        return " ".join(overlap_words) + " "
    
    def _extract_chunk_metadata(self, chunk_text: str) -> Dict[str, Any]:
        """Extract metadata specific to a chunk."""
        metadata = {}
        
        # Contar entidades jurídicas no chunk
        for entity_type, pattern in self.legal_patterns.items():
            matches = pattern.findall(chunk_text)
            if matches:
                metadata[f'{entity_type}_count'] = len(matches)
                metadata[f'{entity_type}_mentions'] = matches[:3]  # Primeiras 3 menções
        
        # Estatísticas básicas
        metadata['word_count'] = len(chunk_text.split())
        metadata['char_count'] = len(chunk_text)
        
        return metadata
    
    def _calculate_hash(self, text: str) -> str:
        """Calculate hash of document content."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    async def process_multiple_documents(self, file_paths: List[str], metadata_list: List[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Process multiple documents concurrently."""
        if metadata_list is None:
            metadata_list = [{}] * len(file_paths)
        
        tasks = []
        for i, file_path in enumerate(file_paths):
            metadata = metadata_list[i] if i < len(metadata_list) else {}
            task = self.process_document(file_path, metadata)
            tasks.append(task)
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filtrar erros
        successful_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Erro ao processar {file_paths[i]}: {str(result)}")
            else:
                successful_results.append(result)
        
        logger.info(f"Processados {len(successful_results)}/{len(file_paths)} documentos com sucesso")
        return successful_results
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_formats)
    
    def validate_document(self, file_path: str) -> bool:
        """Validate if document can be processed."""
        try:
            file_path = Path(file_path)
            return (
                file_path.exists() and
                file_path.is_file() and
                file_path.suffix.lower() in self.supported_formats and
                file_path.stat().st_size > 0
            )
        except:
            return False