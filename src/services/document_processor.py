"""
Document Processor para JurisOracle
Processador avançado de documentos jurídicos com suporte a múltiplos formatos,
extração de metadados, chunking inteligente e integração com pipeline RAG
"""

import os
import re
import io
import json
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union, BinaryIO
from dataclasses import dataclass, asdict
from datetime import datetime
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
from abc import ABC, abstractmethod

# Bibliotecas para processamento de documentos
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
import openpyxl
from odf import text, teletype
from odf.opendocument import load as load_odt

# Bibliotecas para análise de texto
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy

# Bibliotecas para processamento de imagens em PDFs
from PIL import Image
import fitz  # PyMuPDF para melhor extração de PDF

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download de recursos NLTK necessários
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

@dataclass
class DocumentMetadata:
    """Metadados extraídos de um documento"""
    filename: str
    file_path: str
    file_size: int
    file_type: str
    mime_type: str
    creation_date: Optional[datetime]
    modification_date: Optional[datetime]
    author: Optional[str]
    title: Optional[str]
    subject: Optional[str]
    creator: Optional[str]
    producer: Optional[str]
    pages: int
    word_count: int
    character_count: int
    language: str
    encoding: str
    hash_md5: str
    hash_sha256: str
    extracted_at: datetime
    
    # Metadados específicos para documentos jurídicos
    document_type: Optional[str] = None  # sentença, acórdão, petição, etc.
    court: Optional[str] = None
    case_number: Optional[str] = None
    parties: List[str] = None
    legal_topics: List[str] = None
    citations: List[str] = None
    
    def __post_init__(self):
        if self.parties is None:
            self.parties = []
        if self.legal_topics is None:
            self.legal_topics = []
        if self.citations is None:
            self.citations = []

@dataclass
class TextChunk:
    """Chunk de texto extraído de documento"""
    id: str
    content: str
    start_page: int
    end_page: int
    start_position: int
    end_position: int
    chunk_type: str  # paragraph, table, header, footer, etc.
    formatting: Dict[str, Any]
    metadata: Dict[str, Any]
    parent_section: Optional[str] = None
    
@dataclass
class ProcessedDocument:
    """Documento processado completo"""
    metadata: DocumentMetadata
    text_content: str
    chunks: List[TextChunk]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    structure: Dict[str, Any]
    processing_stats: Dict[str, Any]
    errors: List[str]
    warnings: List[str]

class DocumentExtractor(ABC):
    """Interface base para extratores de documento"""
    
    @abstractmethod
    def can_process(self, file_path: str, mime_type: str) -> bool:
        """Verifica se pode processar o tipo de arquivo"""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extrai metadados do documento"""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extrai texto do documento"""
        pass
    
    @abstractmethod
    def extract_structured_content(self, file_path: str) -> Tuple[str, List[TextChunk]]:
        """Extrai conteúdo estruturado com chunks"""
        pass

class PDFExtractor(DocumentExtractor):
    """Extrator especializado para arquivos PDF"""
    
    def __init__(self):
        self.supported_types = ['application/pdf', '.pdf']
    
    def can_process(self, file_path: str, mime_type: str) -> bool:
        return (mime_type in self.supported_types or 
                file_path.lower().endswith('.pdf'))
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extrai metadados de PDF"""
        try:
            file_stats = os.stat(file_path)
            
            # Usa PyMuPDF para metadados mais completos
            doc = fitz.open(file_path)
            pdf_metadata = doc.metadata
            
            # Calcula hashes
            with open(file_path, 'rb') as f:
                content = f.read()
                md5_hash = hashlib.md5(content).hexdigest()
                sha256_hash = hashlib.sha256(content).hexdigest()
            
            # Extrai texto para contagem
            text_content = self.extract_text(file_path)
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            # Detecta idioma (simplificado)
            language = self._detect_language(text_content)
            
            # Extrai metadados jurídicos específicos
            legal_metadata = self._extract_legal_metadata(text_content)
            
            return DocumentMetadata(
                filename=os.path.basename(file_path),
                file_path=file_path,
                file_size=file_stats.st_size,
                file_type='PDF',
                mime_type='application/pdf',
                creation_date=datetime.fromtimestamp(file_stats.st_ctime),
                modification_date=datetime.fromtimestamp(file_stats.st_mtime),
                author=pdf_metadata.get('author'),
                title=pdf_metadata.get('title'),
                subject=pdf_metadata.get('subject'),
                creator=pdf_metadata.get('creator'),
                producer=pdf_metadata.get('producer'),
                pages=doc.page_count,
                word_count=word_count,
                character_count=char_count,
                language=language,
                encoding='UTF-8',
                hash_md5=md5_hash,
                hash_sha256=sha256_hash,
                extracted_at=datetime.now(),
                **legal_metadata
            )
            
        except Exception as e:
            logger.error(f"Erro ao extrair metadados PDF: {e}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto simples do PDF"""
        text_content = ""
        
        try:
            # Primeiro tenta com PyMuPDF (melhor para texto)
            doc = fitz.open(file_path)
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_content += page.get_text()
                text_content += "\n\n"
            
            doc.close()
            
            # Se não conseguiu extrair texto, tenta com pdfplumber
            if not text_content.strip():
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
            
            return self._clean_extracted_text(text_content)
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto PDF: {e}")
            raise
    
    def extract_structured_content(self, file_path: str) -> Tuple[str, List[TextChunk]]:
        """Extrai conteúdo estruturado do PDF"""
        chunks = []
        full_text = ""
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                
                # Extrai blocos de texto com posicionamento
                blocks = page.get_text("dict")["blocks"]
                
                for block_idx, block in enumerate(blocks):
                    if "lines" in block:  # Bloco de texto
                        block_text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                block_text += span["text"]
                        
                        if block_text.strip():
                            chunk_id = f"page_{page_num}_block_{block_idx}"
                            
                            # Determina tipo do chunk baseado na formatação
                            chunk_type = self._determine_chunk_type(block, block_text)
                            
                            # Extrai informações de formatação
                            formatting = self._extract_formatting_info(block)
                            
                            chunk = TextChunk(
                                id=chunk_id,
                                content=block_text.strip(),
                                start_page=page_num + 1,
                                end_page=page_num + 1,
                                start_position=len(full_text),
                                end_position=len(full_text) + len(block_text),
                                chunk_type=chunk_type,
                                formatting=formatting,
                                metadata={
                                    "bbox": block["bbox"],
                                    "block_index": block_idx
                                }
                            )
                            
                            chunks.append(chunk)
                            full_text += block_text + "\n"
            
            doc.close()
            return self._clean_extracted_text(full_text), chunks
            
        except Exception as e:
            logger.error(f"Erro ao extrair conteúdo estruturado PDF: {e}")
            # Fallback para extração simples
            simple_text = self.extract_text(file_path)
            simple_chunks = self._create_simple_chunks(simple_text, "pdf")
            return simple_text, simple_chunks
    
    def _determine_chunk_type(self, block: Dict, text: str) -> str:
        """Determina o tipo de chunk baseado na formatação e conteúdo"""
        if not text.strip():
            return "empty"
        
        # Verifica se é título/cabeçalho (fonte maior, negrito)
        if any(span.get("flags", 0) & 2**4 for line in block.get("lines", []) 
               for span in line.get("spans", [])):  # Bold flag
            return "header"
        
        # Verifica se é número de página
        if re.match(r'^\s*\d+\s*$', text):
            return "page_number"
        
        # Verifica se é citação legal
        if re.search(r'art\.?\s*\d+|lei\s+n[°º]?\s*\d+|cf\.|vide|§', text.lower()):
            return "legal_citation"
        
        # Verifica se parece com tabela
        if '\t' in text or len(re.findall(r'\s{3,}', text)) > 2:
            return "table_content"
        
        return "paragraph"
    
    def _extract_formatting_info(self, block: Dict) -> Dict[str, Any]:
        """Extrai informações de formatação do bloco"""
        formatting = {
            "bold": False,
            "italic": False,
            "font_size": 12,
            "font_family": "Unknown"
        }
        
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                flags = span.get("flags", 0)
                formatting["bold"] = bool(flags & 2**4)  # Bold
                formatting["italic"] = bool(flags & 2**1)  # Italic
                formatting["font_size"] = span.get("size", 12)
                formatting["font_family"] = span.get("font", "Unknown")
                break  # Usa formatação do primeiro span
            if formatting["font_family"] != "Unknown":
                break
        
        return formatting

class DOCXExtractor(DocumentExtractor):
    """Extrator especializado para arquivos DOCX"""
    
    def __init__(self):
        self.supported_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.docx'
        ]
    
    def can_process(self, file_path: str, mime_type: str) -> bool:
        return (mime_type in self.supported_types or 
                file_path.lower().endswith('.docx'))
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extrai metadados de DOCX"""
        try:
            file_stats = os.stat(file_path)
            doc = DocxDocument(file_path)
            
            # Calcula hashes
            with open(file_path, 'rb') as f:
                content = f.read()
                md5_hash = hashlib.md5(content).hexdigest()
                sha256_hash = hashlib.sha256(content).hexdigest()
            
            # Extrai texto para contagem
            text_content = self.extract_text(file_path)
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            # Conta páginas (aproximado)
            pages = max(1, len(text_content) // 2500)  # Estimativa
            
            # Detecta idioma
            language = self._detect_language(text_content)
            
            # Metadados do documento
            core_props = doc.core_properties
            
            # Extrai metadados jurídicos
            legal_metadata = self._extract_legal_metadata(text_content)
            
            return DocumentMetadata(
                filename=os.path.basename(file_path),
                file_path=file_path,
                file_size=file_stats.st_size,
                file_type='DOCX',
                mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                creation_date=core_props.created,
                modification_date=core_props.modified,
                author=core_props.author,
                title=core_props.title,
                subject=core_props.subject,
                creator=core_props.author,
                producer='Microsoft Word',
                pages=pages,
                word_count=word_count,
                character_count=char_count,
                language=language,
                encoding='UTF-8',
                hash_md5=md5_hash,
                hash_sha256=sha256_hash,
                extracted_at=datetime.now(),
                **legal_metadata
            )
            
        except Exception as e:
            logger.error(f"Erro ao extrair metadados DOCX: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto simples do DOCX"""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extrai texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            return self._clean_extracted_text(text_content)
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto DOCX: {e}")
            raise
    
    def extract_structured_content(self, file_path: str) -> Tuple[str, List[TextChunk]]:
        """Extrai conteúdo estruturado do DOCX"""
        chunks = []
        full_text = ""
        
        try:
            doc = DocxDocument(file_path)
            
            # Processa parágrafos
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    chunk_id = f"paragraph_{para_idx}"
                    
                    # Determina tipo baseado no estilo
                    chunk_type = self._determine_docx_chunk_type(paragraph)
                    
                    # Extrai formatação
                    formatting = self._extract_docx_formatting(paragraph)
                    
                    chunk = TextChunk(
                        id=chunk_id,
                        content=paragraph.text.strip(),
                        start_page=1,  # DOCX não tem páginas fixas
                        end_page=1,
                        start_position=len(full_text),
                        end_position=len(full_text) + len(paragraph.text),
                        chunk_type=chunk_type,
                        formatting=formatting,
                        metadata={
                            "style": paragraph.style.name if paragraph.style else "Normal",
                            "paragraph_index": para_idx
                        }
                    )
                    
                    chunks.append(chunk)
                    full_text += paragraph.text + "\n"
            
            # Processa tabelas
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    chunk_id = f"table_{table_idx}"
                    
                    chunk = TextChunk(
                        id=chunk_id,
                        content=table_text,
                        start_page=1,
                        end_page=1,
                        start_position=len(full_text),
                        end_position=len(full_text) + len(table_text),
                        chunk_type="table",
                        formatting={"table_rows": len(table.rows), "table_cols": len(table.columns)},
                        metadata={"table_index": table_idx}
                    )
                    
                    chunks.append(chunk)
                    full_text += table_text + "\n"
            
            return self._clean_extracted_text(full_text), chunks
            
        except Exception as e:
            logger.error(f"Erro ao extrair conteúdo estruturado DOCX: {e}")
            simple_text = self.extract_text(file_path)
            simple_chunks = self._create_simple_chunks(simple_text, "docx")
            return simple_text, simple_chunks
    
    def _determine_docx_chunk_type(self, paragraph) -> str:
        """Determina tipo do chunk DOCX baseado no estilo"""
        style_name = paragraph.style.name.lower() if paragraph.style else "normal"
        
        if "heading" in style_name or "título" in style_name:
            return "header"
        elif "quote" in style_name or "citação" in style_name:
            return "quote"
        elif "caption" in style_name or "legenda" in style_name:
            return "caption"
        else:
            return "paragraph"
    
    def _extract_docx_formatting(self, paragraph) -> Dict[str, Any]:
        """Extrai formatação do parágrafo DOCX"""
        formatting = {
            "bold": False,
            "italic": False,
            "underline": False,
            "font_size": 12,
            "font_family": "Calibri",
            "alignment": "left"
        }
        
        if paragraph.runs:
            # Usa formatação do primeiro run
            run = paragraph.runs[0]
            formatting["bold"] = run.bold or False
            formatting["italic"] = run.italic or False
            formatting["underline"] = run.underline or False
            
            if run.font.size:
                formatting["font_size"] = run.font.size.pt
            if run.font.name:
                formatting["font_family"] = run.font.name
        
        # Alinhamento do parágrafo
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
        }
        formatting["alignment"] = alignment_map.get(paragraph.alignment, "left")
        
        return formatting
    
    def _extract_table_text(self, table) -> str:
        """Extrai texto formatado de tabela DOCX"""
        table_text = ""
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                row_text.append(cell_text)
            table_text += " | ".join(row_text) + "\n"
        return table_text

class DocumentProcessor:
    """
    Processador principal de documentos para JurisOracle
    
    Funcionalidades:
    - Suporte a múltiplos formatos (PDF, DOCX, ODT, TXT)
    - Extração de metadados avançada
    - Chunking inteligente com preservação de estrutura
    - Processamento assíncrono para grandes volumes
    - Cache de documentos processados
    - Análise de conteúdo jurídico
    - Detecção de duplicatas
    """
    
    def __init__(
        self,
        cache_dir: str = "./data/document_cache",
        max_workers: int = 4,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        enable_nlp: bool = True
    ):
        """
        Inicializa o processador de documentos
        
        Args:
            cache_dir: Diretório para cache de documentos
            max_workers: Número máximo de workers para processamento paralelo
            chunk_size: Tamanho padrão dos chunks
            chunk_overlap: Sobreposição entre chunks
            enable_nlp: Habilita processamento NLP avançado
        """
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.max_workers = max_workers
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_nlp = enable_nlp
        
        # Inicializa extratores
        self.extractors = [
            PDFExtractor(),
            DOCXExtractor()
        ]
        
        # Inicializa NLP se habilitado
        if self.enable_nlp:
            self._initialize_nlp()
        
        # Estatísticas
        self.processing_stats = {
            "documents_processed": 0,
            "total_processing_time": 0.0,
            "errors": 0,
            "cache_hits": 0
        }
        
        logger.info("Document Processor inicializado")
    
    def _initialize_nlp(self):
        """Inicializa componentes NLP"""
        try:
            # Carrega modelo spaCy para português
            self.nlp = spacy.load("pt_core_news_sm")
            logger.info("Modelo spaCy carregado para português")
        except OSError:
            logger.warning("Modelo spaCy pt_core_news_sm não encontrado. Usando processamento básico.")
            self.nlp = None
        
        # Stopwords em português
        try:
            self.stop_words = set(stopwords.words('portuguese'))
        except:
            self.stop_words = set()
    
    def _get_file_extractor(self, file_path: str) -> Optional[DocumentExtractor]:
        """Encontra extrator apropriado para o arquivo"""
        mime_type = mimetypes.guess_type(file_path)[0] or ""
        
        for extractor in self.extractors:
            if extractor.can_process(file_path, mime_type):
                return extractor
        
        return None
    
    def _generate_cache_key(self, file_path: str) -> str:
        """Gera chave de cache baseada no arquivo"""
        file_stats = os.stat(file_path)
        content = f"{file_path}_{file_stats.st_size}_{file_stats.st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def _save_to_cache(self, cache_key: str, processed_doc: ProcessedDocument):
        """Salva documento processado no cache"""
        try:
            cache_file = self.cache_dir / f"{cache_key}.json"
            
            # Converte para formato serializável
            cache_data = {
                "metadata": asdict(processed_doc.metadata),
                "text_content": processed_doc.text_content,
                "chunks": [asdict(chunk) for chunk in processed_doc.chunks],
                "tables": processed_doc.tables,
                "images": processed_doc.images,
                "structure": processed_doc.structure,
                "processing_stats": processed_doc.processing_stats,
                "errors": processed_doc.errors,
                "warnings": processed_doc.warnings
            }
            
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2, default=str)
                
        except Exception as e:
            logger.warning(f"Erro ao salvar cache: {e}")
    
    def _load_from_cache(self, cache_key: str) -> Optional[ProcessedDocument]:
        """Carrega documento do cache"""
        try:
            cache_file = self.cache_dir / f"{cache_key}.json"
            
            if not cache_file.exists():
                return None
            
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            # Reconstrói objetos
            metadata = DocumentMetadata(**cache_data["metadata"])
            chunks = [TextChunk(**chunk_data) for chunk_data in cache_data["chunks"]]
            
            processed_doc = ProcessedDocument(
                metadata=metadata,
                text_content=cache_data["text_content"],
                chunks=chunks,
                tables=cache_data["tables"],
                images=cache_data["images"],
                structure=cache_data["structure"],
                processing_stats=cache_data["processing_stats"],
                errors=cache_data["errors"],
                warnings=cache_data["warnings"]
            )
            
            self.processing_stats["cache_hits"] += 1
            return processed_doc
            
        except Exception as e:
            logger.warning(f"Erro ao carregar cache: {e}")
            return None
    
    def process_document(
        self,
        file_path: str,
        use_cache: bool = True,
        extract_images: bool = False,
        extract_tables: bool = True
    ) -> ProcessedDocument:
        """
        Processa um único documento
        
        Args:
            file_path: Caminho para o arquivo
            use_cache: Usar cache se disponível
            extract_images: Extrair informações de imagens
            extract_tables: Extrair informações de tabelas
            
        Returns:
            Documento processado
        """
        start_time = datetime.now()
        
        try:
            # Verifica cache
            cache_key = self._generate_cache_key(file_path)
            if use_cache:
                cached_doc = self._load_from_cache(cache_key)
                if cached_doc:
                    logger.info(f"Documento carregado do cache: {file_path}")
                    return cached_doc
            
            # Encontra extrator
            extractor = self._get_file_extractor(file_path)
            if not extractor:
                raise ValueError(f"Formato de arquivo não suportado: {file_path}")
            
            logger.info(f"Processando documento: {file_path}")
            
            # Extrai metadados
            metadata = extractor.extract_metadata(file_path)
            
            # Extrai conteúdo estruturado
            text_content, chunks = extractor.extract_structured_content(file_path)
            
            # Processa chunks adicionais se necessário
            if len(chunks) == 0 or all(len(chunk.content) > self.chunk_size for chunk in chunks):
                additional_chunks = self._create_intelligent_chunks(text_content, metadata.file_type)
                chunks.extend(additional_chunks)
            
            # Extrai tabelas se solicitado
            tables = []
            if extract_tables:
                tables = self._extract_tables(file_path, extractor)
            
            # Extrai imagens se solicitado
            images = []
            if extract_images:
                images = self._extract_images(file_path, extractor)
            
            # Analisa estrutura do documento
            structure = self._analyze_document_structure(chunks, metadata)
            
            # Processa NLP se habilitado
            if self.enable_nlp and self.nlp:
                self._enrich_with_nlp(chunks, text_content)
            
            # Estatísticas de processamento
            processing_time = (datetime.now() - start_time).total_seconds()
            processing_stats = {
                "processing_time": processing_time,
                "extractor_used": extractor.__class__.__name__,
                "chunks_created": len(chunks),
                "tables_found": len(tables),
                "images_found": len(images)
            }
            
            # Cria documento processado
            processed_doc = ProcessedDocument(
                metadata=metadata,
                text_content=text_content,
                chunks=chunks,
                tables=tables,
                images=images,
                structure=structure,
                processing_stats=processing_stats,
                errors=[],
                warnings=[]
            )
            
            # Salva no cache
            if use_cache:
                self._save_to_cache(cache_key, processed_doc)
            
            # Atualiza estatísticas
            self.processing_stats["documents_processed"] += 1
            self.processing_stats["total_processing_time"] += processing_time
            
            logger.info(f"Documento processado em {processing_time:.2f}s: {len(chunks)} chunks criados")
            
            return processed_doc
            
        except Exception as e:
            self.processing_stats["errors"] += 1
            logger.error(f"Erro ao processar documento {file_path}: {e}")
            raise
    
    def process_documents_batch(
        self,
        file_paths: List[str],
        use_cache: bool = True,
        extract_images: bool = False,
        extract_tables: bool = True
    ) -> List[ProcessedDocument]:
        """
        Processa múltiplos documentos em paralelo
        
        Args:
            file_paths: Lista de caminhos de arquivos
            use_cache: Usar cache se disponível
            extract_images: Extrair informações de imagens
            extract_tables: Extrair informações de tabelas
            
        Returns:
            Lista de documentos processados
        """
        logger.info(f"Processando batch de {len(file_paths)} documentos")
        
        processed_docs = []
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submete tarefas
            future_to_file = {
                executor.submit(
                    self.process_document,
                    file_path,
                    use_cache,
                    extract_images,
                    extract_tables
                ): file_path for file_path in file_paths
            }
            
            # Coleta resultados
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    processed_doc = future.result()
                    processed_docs.append(processed_doc)
                except Exception as e:
                    logger.error(f"Erro ao processar {file_path}: {e}")
                    # Cria documento de erro
                    error_doc = self._create_error_document(file_path, str(e))
                    processed_docs.append(error_doc)
        
        return processed_docs
    
    def _create_intelligent_chunks(
        self,
        text: str,
        file_type: str
    ) -> List[TextChunk]:
        """Cria chunks inteligentes baseados no conteúdo"""
        chunks = []
        
        # Tokeniza em sentenças
        sentences = sent_tokenize(text, language='portuguese')
        
        current_chunk = ""
        current_chunk_sentences = []
        chunk_start = 0
        
        for i, sentence in enumerate(sentences):
            # Verifica se adicionar a sentença excederia o tamanho
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Cria chunk atual
                chunk_id = f"intelligent_chunk_{len(chunks)}"
                chunk = TextChunk(
                    id=chunk_id,
                    content=current_chunk.strip(),
                    start_page=1,
                    end_page=1,
                    start_position=chunk_start,
                    end_position=chunk_start + len(current_chunk),
                    chunk_type="intelligent_paragraph",
                    formatting={},
                    metadata={
                        "sentence_count": len(current_chunk_sentences),
                        "file_type": file_type
                    }
                )
                chunks.append(chunk)
                
                # Inicia novo chunk com overlap
                overlap_sentences = current_chunk_sentences[-self.chunk_overlap // 100:]
                current_chunk = " ".join(overlap_sentences) + " " + sentence
                current_chunk_sentences = overlap_sentences + [sentence]
                chunk_start += len(current_chunk) - len(" ".join(overlap_sentences)) - len(sentence)
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_chunk_sentences.append(sentence)
        
        # Adiciona último chunk
        if current_chunk.strip():
            chunk_id = f"intelligent_chunk_{len(chunks)}"
            chunk = TextChunk(
                id=chunk_id,
                content=current_chunk.strip(),
                start_page=1,
                end_page=1,
                start_position=chunk_start,
                end_position=chunk_start + len(current_chunk),
                chunk_type="intelligent_paragraph",
                formatting={},
                metadata={
                    "sentence_count": len(current_chunk_sentences),
                    "file_type": file_type
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def _extract_tables(self, file_path: str, extractor: DocumentExtractor) -> List[Dict[str, Any]]:
        """Extrai informações de tabelas (implementação específica por extrator)"""
        # Implementação básica - pode ser estendida
        return []
    
    def _extract_images(self, file_path: str, extractor: DocumentExtractor) -> List[Dict[str, Any]]:
        """Extrai informações de imagens (implementação específica por extrator)"""
        # Implementação básica - pode ser estendida
        return []
    
    def _analyze_document_structure(
        self,
        chunks: List[TextChunk],
        metadata: DocumentMetadata
    ) -> Dict[str, Any]:
        """Analisa a estrutura do documento"""
        structure = {
            "has_headers": any(chunk.chunk_type == "header" for chunk in chunks),
            "has_tables": any(chunk.chunk_type == "table" for chunk in chunks),
            "has_legal_citations": any(chunk.chunk_type == "legal_citation" for chunk in chunks),
            "section_count": len([c for c in chunks if c.chunk_type == "header"]),
            "paragraph_count": len([c for c in chunks if c.chunk_type == "paragraph"]),
            "average_chunk_length": np.mean([len(c.content) for c in chunks]) if chunks else 0,
            "document_type": self._classify_document_type(chunks, metadata)
        }
        
        return structure
    
    def _classify_document_type(
        self,
        chunks: List[TextChunk],
        metadata: DocumentMetadata
    ) -> str:
        """Classifica o tipo de documento jurídico"""
        text_sample = " ".join([chunk.content for chunk in chunks[:5]]).lower()
        
        # Padrões para diferentes tipos de documentos
        patterns = {
            "sentenca": ["sentença", "julgo", "dispositivo", "condeno"],
            "acordao": ["acórdão", "turma", "relator", "revisor"],
            "peticao": ["petição", "requer", "solicita", "postula"],
            "contrato": ["contrato", "partes", "cláusula", "contratante"],
            "lei": ["lei", "decreto", "artigo", "parágrafo"],
            "parecer": ["parecer", "opinião", "consulta", "entendimento"]
        }
        
        for doc_type, keywords in patterns.items():
            if sum(1 for keyword in keywords if keyword in text_sample) >= 2:
                return doc_type
        
        return "documento_generico"
    
    def _enrich_with_nlp(self, chunks: List[TextChunk], text_content: str):
        """Enriquece chunks com análise NLP"""
        if not self.nlp:
            return
        
        # Análise do documento completo
        doc = self.nlp(text_content[:1000000])  # Limita tamanho para performance
        
        # Extrai entidades nomeadas
        entities = [(ent.text, ent.label_) for ent in doc.ents]
        
        # Enriquece chunks individuais
        for chunk in chunks:
            if len(chunk.content) < 10000:  # Processa apenas chunks menores
                chunk_doc = self.nlp(chunk.content)
                
                # Adiciona informações NLP aos metadados
                chunk.metadata.update({
                    "entities": [(ent.text, ent.label_) for ent in chunk_doc.ents],
                    "pos_tags": [(token.text, token.pos_) for token in chunk_doc[:20]],  # Primeiros 20 tokens
                    "sentiment": "neutral"  # Placeholder para análise de sentimento
                })
    
    def _create_error_document(self, file_path: str, error: str) -> ProcessedDocument:
        """Cria documento de erro para falhas de processamento"""
        error_metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_path=file_path,
            file_size=0,
            file_type="ERROR",
            mime_type="application/error",
            creation_date=datetime.now(),
            modification_date=datetime.now(),
            author=None,
            title=None,
            subject=None,
            creator=None,
            producer=None,
            pages=0,
            word_count=0,
            character_count=0,
            language="unknown",
            encoding="unknown",
            hash_md5="",
            hash_sha256="",
            extracted_at=datetime.now()
        )
        
        return ProcessedDocument(
            metadata=error_metadata,
            text_content="",
            chunks=[],
            tables=[],
            images=[],
            structure={},
            processing_stats={"processing_time": 0.0},
            errors=[error],
            warnings=[]
        )
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Retorna estatísticas de processamento"""
        return self.processing_stats.copy()
    
    def clear_cache(self):
        """Limpa cache de documentos"""
        try:
            for cache_file in self.cache_dir.glob("*.json"):
                cache_file.unlink()
            logger.info("Cache limpo com sucesso")
        except Exception as e:
            logger.error(f"Erro ao limpar cache: {e}")
    
    def _detect_language(self, text: str) -> str:
        """Detecta idioma do texto (implementação básica)"""
        # Palavras comuns em português
        portuguese_words = ['de', 'da', 'do', 'que', 'e', 'a', 'o', 'para', 'com', 'não']
        text_sample = text.lower()[:1000]
        
        pt_count = sum(1 for word in portuguese_words if word in text_sample)
        return "portuguese" if pt_count >= 3 else "unknown"
    
    def _extract_legal_metadata(self, text: str) -> Dict[str, Any]:
        """Extrai metadados específicos de documentos jurídicos"""
        legal_metadata = {
            "document_type": None,
            "court": None,
            "case_number": None,
            "parties": [],
            "legal_topics": [],
            "citations": []
        }
        
        # Padrões regex para extração
        case_number_pattern = r'\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}'
        court_patterns = [
            r'Supremo Tribunal Federal',
            r'Superior Tribunal de Justiça',
            r'Tribunal de Justiça',
            r'Tribunal Regional Federal'
        ]
        
        # Extrai número do processo
        case_matches = re.findall(case_number_pattern, text)
        if case_matches:
            legal_metadata["case_number"] = case_matches[0]
        
        # Extrai tribunal
        for pattern in court_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                legal_metadata["court"] = pattern
                break
        
        # Extrai citações legais (simplificado)
        citation_patterns = [
            r'[Aa]rt\.?\s*\d+',
            r'[Ll]ei\s+n[°º]?\s*\d+',
            r'CF/88',
            r'CC/02'
        ]
        
        citations = []
        for pattern in citation_patterns:
            citations.extend(re.findall(pattern, text))
        legal_metadata["citations"] = citations
        
        return legal_metadata
    
    def _clean_extracted_text(self, text: str) -> str:
        """Limpa texto extraído removendo artefatos"""
        # Remove quebras de linha excessivas
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove espaços extras
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove caracteres de controle
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _create_simple_chunks(self, text: str, file_type: str) -> List[TextChunk]:
        """Cria chunks simples quando a extração estruturada falha"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text.strip()) < 50:
                continue
            
            chunk_id = f"simple_chunk_{len(chunks)}"
            chunk = TextChunk(
                id=chunk_id,
                content=chunk_text,
                start_page=1,
                end_page=1,
                start_position=i,
                end_position=i + len(chunk_words),
                chunk_type="simple_paragraph",
                formatting={},
                metadata={"file_type": file_type, "method": "simple_chunking"}
            )
            
            chunks.append(chunk)
        
        return chunks

# Utilitários adicionais para integração com RAG
def integrate_with_rag_pipeline(
    document_processor: DocumentProcessor,
    rag_pipeline,
    file_paths: List[str]
) -> Dict[str, Any]:
    """
    Integra processamento de documentos com pipeline RAG
    
    Args:
        document_processor: Instância do processador
        rag_pipeline: Pipeline RAG para adicionar documentos
        file_paths: Lista de arquivos para processar
        
    Returns:
        Estatísticas de integração
    """
    logger.info(f"Integrando {len(file_paths)} documentos com pipeline RAG")
    
    # Processa documentos
    processed_docs = document_processor.process_documents_batch(file_paths)
    
    # Adiciona ao pipeline RAG
    total_chunks = 0
    successful_docs = 0
    
    for doc in processed_docs:
        if doc.errors:
            logger.warning(f"Documento com erros ignorado: {doc.metadata.filename}")
            continue
        
        try:
            # Adiciona documento ao RAG
            chunks_added = rag_pipeline.add_document(
                content=doc.text_content,
                document_id=doc.metadata.filename,
                metadata=asdict(doc.metadata)
            )
            
            total_chunks += chunks_added
            successful_docs += 1
            
        except Exception as e:
            logger.error(f"Erro ao adicionar documento ao RAG: {e}")
    
    # Salva pipeline RAG
    rag_pipeline.save_vector_store()
    
    return {
        "total_documents": len(file_paths),
        "processed_documents": len(processed_docs),
        "successful_integrations": successful_docs,
        "total_chunks_added": total_chunks,
        "errors": len([doc for doc in processed_docs if doc.errors])
    }